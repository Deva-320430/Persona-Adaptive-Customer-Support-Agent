"""
Ingestion pipeline: loads documents from data/, chunks them,
embeds them locally, and builds a FAISS index.

Run this once (or whenever data/ changes):
    python src/ingestion.py
"""
import os
import json
import sys

import faiss
import numpy as np
from pypdf import PdfReader
from docx import Document

import config
import llm_client


def load_text_file(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_pdf_file(path):
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def load_docx_file(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables are common in policy/FAQ docs (e.g. pricing tiers) — without this,
    # any knowledge stored in a table would silently disappear from the index.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_documents(data_dir):
    """Returns a list of {"text": ..., "source": filename}."""
    docs = []
    if not os.path.isdir(data_dir):
        raise FileNotFoundError(f"Data directory not found: {data_dir}")

    for filename in sorted(os.listdir(data_dir)):
        path = os.path.join(data_dir, filename)
        if not os.path.isfile(path):
            continue

        if filename.lower().endswith(".pdf"):
            text = load_pdf_file(path)
        elif filename.lower().endswith(".docx"):
            text = load_docx_file(path)
        elif filename.lower().endswith(".doc"):
            print(f"  [skipped] {filename}: old binary .doc format isn't supported, "
                  f"only .docx. Re-save it as .docx in Word and re-run ingestion.")
            continue
        elif filename.lower().endswith((".md", ".txt")):
            text = load_text_file(path)
        else:
            # Loud warning instead of silently skipping — a forgotten .csv or
            # image dropped into data/ should be obvious, not a silent gap
            # in the knowledge base that's hard to debug later.
            print(f"  [skipped] {filename}: unsupported file type (supported: .pdf, .docx, .md, .txt)")
            continue

        if text.strip():
            docs.append({"text": text, "source": filename})

    return docs


def chunk_text(text, chunk_size=500, overlap=100):
    """
    Simple sliding-window chunker over characters. Good enough for a
    knowledge base of short markdown docs; for longer/denser docs a
    sentence-aware splitter would do better, but this keeps the
    pipeline dependency-light and easy to reason about.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    step = max(chunk_size - overlap, 1)  # guard against overlap >= chunk_size
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += step
    return chunks


def build_index(data_dir=None, output_dir=None):
    data_dir = data_dir or config.DATA_DIR
    output_dir = output_dir or config.VECTORSTORE_DIR
    os.makedirs(output_dir, exist_ok=True)

    print(f"Loading documents from {data_dir} ...")
    docs = load_documents(data_dir)
    if not docs:
        raise ValueError(f"No documents found in {data_dir}. Add .md, .txt, or .pdf files.")
    print(f"Loaded {len(docs)} documents.")

    chunks = []  # list of {"text", "source", "chunk_id"}
    for doc in docs:
        text_chunks = chunk_text(doc["text"], config.CHUNK_SIZE, config.CHUNK_OVERLAP)
        for i, chunk in enumerate(text_chunks):
            chunks.append({
                "text": chunk,
                "source": doc["source"],
                "chunk_id": i,
            })
    print(f"Split into {len(chunks)} chunks.")

    texts = [c["text"] for c in chunks]
    print(f"Embedding {len(texts)} chunks via Gemini ('{config.EMBEDDING_MODEL}') ...")
    embeddings_list = llm_client.embed_texts(texts, task_type="RETRIEVAL_DOCUMENT")
    embeddings = np.array(embeddings_list, dtype=np.float32)

    # Normalize so inner product == cosine similarity, and HIGHER score == better match.
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]

    index = faiss.IndexFlatIP(dim)  # inner product on normalized vectors = cosine sim
    index.add(embeddings.astype(np.float32))

    faiss.write_index(index, config.INDEX_PATH)
    with open(config.METADATA_PATH, "w", encoding="utf-8") as f:
        json.dump(chunks, f, indent=2)

    print(f"Saved FAISS index to {config.INDEX_PATH}")
    print(f"Saved metadata to {config.METADATA_PATH}")
    print("Ingestion complete.")


if __name__ == "__main__":
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    build_index()
